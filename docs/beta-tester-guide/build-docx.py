from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
MARKDOWN_PATH = ROOT / "README.md"
DOCX_PATH = ROOT / "capitol-ledger-beta-tester-guide.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 31, 43)
MUTED = RGBColor(82, 94, 112)


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_style_font(style, size, color=INK, bold=False):
    font = style.font
    font.name = "Calibri"
    font.size = Pt(size)
    font.color.rgb = color
    font.bold = bold
    style.element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style.element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_text_with_links(paragraph, text):
    url_pattern = re.compile(r"https?://\S+")
    cursor = 0
    for match in url_pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        url = match.group(0)
        add_hyperlink(paragraph, url, url)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def iter_markdown_blocks(markdown):
    paragraph = []
    list_block = None

    def flush_paragraph():
        nonlocal paragraph
        if paragraph:
            yield {"type": "paragraph", "text": " ".join(paragraph)}
            paragraph = []

    def flush_list():
        nonlocal list_block
        if list_block:
            yield list_block
            list_block = None

    for raw in markdown.splitlines():
        line = raw.rstrip()
        image_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", line)
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        bullet_match = re.match(r"^-\s+(.+)$", line)
        number_match = re.match(r"^\d+\.\s+(.+)$", line)

        if image_match:
            yield from flush_paragraph()
            yield from flush_list()
            yield {"type": "image", "alt": image_match.group(1), "src": image_match.group(2)}
        elif heading_match:
            yield from flush_paragraph()
            yield from flush_list()
            yield {"type": f"h{len(heading_match.group(1))}", "text": heading_match.group(2)}
        elif bullet_match:
            yield from flush_paragraph()
            if not list_block or list_block["ordered"]:
                yield from flush_list()
                list_block = {"type": "list", "ordered": False, "items": []}
            list_block["items"].append(bullet_match.group(1))
        elif number_match:
            yield from flush_paragraph()
            if not list_block or not list_block["ordered"]:
                yield from flush_list()
                list_block = {"type": "list", "ordered": True, "items": []}
            list_block["items"].append(number_match.group(1))
        elif not line.strip():
            yield from flush_paragraph()
            yield from flush_list()
        else:
            yield from flush_list()
            paragraph.append(line.strip())

    yield from flush_paragraph()
    yield from flush_list()


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], 11, INK)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        set_style_font(style, size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        set_style_font(style, 11, INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "Capitol Ledger CE Beta Tester Guide"
    set_paragraph_spacing(header, after=0)
    set_run_font(header.runs[0], size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "Editable beta tester handout"
    set_paragraph_spacing(footer, after=0)
    set_run_font(footer.runs[0], size=9, color=MUTED)


def add_title_block(doc):
    title = doc.add_paragraph()
    set_paragraph_spacing(title, after=3)
    run = title.add_run("Capitol Ledger CE First-Round Beta Tester Guide")
    set_run_font(run, size=24, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=14)
    run = subtitle.add_run("Editable version for first-round beta testing")
    set_run_font(run, size=12, color=MUTED, italic=True)


def add_block(doc, block):
    block_type = block["type"]
    if block_type == "h1":
        # The opening H1 is already handled as the document title block.
        return
    if block_type == "h2":
        p = doc.add_paragraph(block["text"], style="Heading 1")
        return
    if block_type == "h3":
        p = doc.add_paragraph(block["text"], style="Heading 2")
        return
    if block_type == "paragraph":
        p = doc.add_paragraph()
        add_text_with_links(p, block["text"])
        set_paragraph_spacing(p)
        return
    if block_type == "list":
        style = "List Number" if block["ordered"] else "List Bullet"
        for item in block["items"]:
            p = doc.add_paragraph(style=style)
            add_text_with_links(p, item)
        return
    if block_type == "image":
        image_path = ROOT / block["src"]
        if not image_path.exists():
            p = doc.add_paragraph(f"[Missing image: {block['alt']}]")
            set_paragraph_spacing(p)
            return
        caption = doc.add_paragraph(block["alt"], style="Heading 3")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph = doc.add_paragraph()
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.add_run().add_picture(str(image_path), width=Inches(5.6))
        set_paragraph_spacing(picture_paragraph, after=10)
        return


def build_docx():
    markdown = MARKDOWN_PATH.read_text()
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Capitol Ledger CE First-Round Beta Tester Guide"
    doc.core_properties.subject = "Editable beta tester guide"
    doc.core_properties.keywords = "Capitol Ledger CE, beta testing, civic app"
    add_title_block(doc)

    for block in iter_markdown_blocks(markdown):
        add_block(doc, block)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
