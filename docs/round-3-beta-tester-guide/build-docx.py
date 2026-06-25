from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor


GUIDE_DIR = Path(__file__).resolve().parent
REPO_ROOT = GUIDE_DIR.parents[1]
SOURCE = GUIDE_DIR / "README.md"
OUTPUT = REPO_ROOT / "public" / "downloads" / "capitol-ledger-round-3-beta-tester-guide.docx"

BLUE = RGBColor(31, 78, 121)
DARK_BLUE = RGBColor(20, 55, 91)
GOLD = RGBColor(181, 121, 16)
TEXT = RGBColor(31, 41, 55)
MUTED = RGBColor(75, 85, 99)
LINK = RGBColor(5, 99, 193)


def set_font(style, size: int, color: RGBColor = TEXT, bold: bool = False) -> None:
    font = style.font
    font.name = "Calibri"
    font.size = Pt(size)
    font.color.rgb = color
    font.bold = bold


def set_spacing(style, before: int = 0, after: int = 6, line: float = 1.25) -> None:
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def ensure_style(document: Document, name: str, base: str = "Normal"):
    styles = document.styles
    if name in styles:
        return styles[name]
    return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def clean_markdown(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    return text


def add_text_or_link(paragraph, text: str) -> None:
    text = clean_markdown(text)
    if text.startswith("http://") or text.startswith("https://"):
        add_hyperlink(paragraph, text, text)
        return
    paragraph.add_run(text)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    set_font(normal, 11, TEXT)
    set_spacing(normal)

    title = document.styles["Title"]
    set_font(title, 18, BLUE, bold=True)
    title.paragraph_format.space_after = Pt(3)

    subtitle = ensure_style(document, "Guide Subtitle")
    set_font(subtitle, 11, MUTED)
    set_spacing(subtitle, after=12)

    h1 = document.styles["Heading 1"]
    set_font(h1, 16, BLUE, bold=True)
    set_spacing(h1, before=12, after=4)

    h2 = document.styles["Heading 2"]
    set_font(h2, 13, DARK_BLUE, bold=True)
    set_spacing(h2, before=8, after=4)

    h3 = document.styles["Heading 3"]
    set_font(h3, 12, DARK_BLUE, bold=True)
    set_spacing(h3, before=8, after=3)

    bullet = document.styles["List Bullet"]
    set_font(bullet, 11, TEXT)
    set_spacing(bullet, after=4)
    bullet.paragraph_format.left_indent = Inches(0.34)
    bullet.paragraph_format.first_line_indent = Inches(-0.18)

    number = document.styles["List Number"]
    set_font(number, 11, TEXT)
    set_spacing(number, after=4)
    number.paragraph_format.left_indent = Inches(0.34)
    number.paragraph_format.first_line_indent = Inches(-0.18)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Capitol Ledger CE Round 3 Beta Tester Guide")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = MUTED


def build_document() -> Document:
    document = Document()
    configure_document(document)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        text = " ".join(paragraph_buffer).strip()
        paragraph_buffer.clear()
        if not text:
            return
        p = document.add_paragraph()
        add_text_or_link(p, text)

    for raw_line in lines:
        line = raw_line.rstrip()

        if not line:
            flush_paragraph()
            continue

        if line.startswith("# "):
            flush_paragraph()
            p = document.add_paragraph(style="Title")
            add_text_or_link(p, line[2:])
            continue

        if line.startswith("Last updated:"):
            flush_paragraph()
            p = document.add_paragraph(style="Guide Subtitle")
            add_text_or_link(p, line)
            continue

        if line.startswith("## "):
            flush_paragraph()
            document.add_paragraph(clean_markdown(line[3:]), style="Heading 1")
            continue

        if line.startswith("### "):
            flush_paragraph()
            document.add_paragraph(clean_markdown(line[4:]), style="Heading 2")
            continue

        if line.startswith("- "):
            flush_paragraph()
            p = document.add_paragraph(style="List Bullet")
            add_text_or_link(p, line[2:])
            continue

        number_match = re.match(r"^\d+\.\s+(.*)$", line)
        if number_match:
            flush_paragraph()
            p = document.add_paragraph(style="List Number")
            add_text_or_link(p, number_match.group(1))
            continue

        paragraph_buffer.append(line)

    flush_paragraph()
    return document


def main() -> None:
    document = build_document()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
